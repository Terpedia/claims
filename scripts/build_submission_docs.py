#!/usr/bin/env python3
"""Build upload-ready Word versions of the title page and cover letter."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SUBMISSION = ROOT / "manuscript" / "submission"


def inline(paragraph, text: str) -> None:
    token = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            run.bold = True
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            run.italic = True
        elif value.startswith("`"):
            paragraph.add_run(value[1:-1])
        else:
            label, _url = re.match(r"\[(.*?)\]\((.*?)\)", value).groups()
            run = paragraph.add_run(label)
            run.underline = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
    title_ppr = doc.styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)


def build(source: Path, output: Path) -> None:
    doc = Document()
    configure(doc)
    for line in source.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            doc.add_paragraph()
        elif line.startswith("# "):
            p = doc.add_paragraph(style="Title")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            inline(p, line[2:])
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            inline(p, line)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


def main() -> None:
    build(SUBMISSION / "title-page.md", SUBMISSION / "title-page.docx")
    build(SUBMISSION / "cover-letter.md", SUBMISSION / "cover-letter.docx")


if __name__ == "__main__":
    main()
