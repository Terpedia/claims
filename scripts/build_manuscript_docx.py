#!/usr/bin/env python3
"""Build a clean Word manuscript from the version-controlled Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "terpene-structure-function-claims.md"
OUTPUT = ROOT / "manuscript" / "submission" / "terpene-structure-function-claims.docx"


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = props.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        props.append(shading)
    shading.set(qn("w:fill"), fill)


def borders(cell, color: str = "D9D9D9") -> None:
    props = cell._tc.get_or_add_tcPr()
    border = props.first_child_found_in("w:tcBorders")
    if border is None:
        border = OxmlElement("w:tcBorders")
        props.append(border)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        node = border.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            border.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), color)


def inline(paragraph, text: str) -> None:
    """Render the small Markdown subset used by the manuscript."""
    token = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            run.bold = True
        elif value.startswith("*"):
            run = paragraph.add_run(value[1:-1])
            run.italic = True
        elif value.startswith("`"):
            paragraph.add_run(value[1:-1])
        else:
            label, url = re.match(r"\[(.*?)\]\((.*?)\)", value).groups()
            run = paragraph.add_run(label)
            run.underline = True
            run.font.color.rgb = RGBColor(31, 78, 121)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_body_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    inline(p, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            borders(cell)
            shade(cell, "1F4E79" if r == 0 else ("F2F6FA" if r % 2 == 0 else "FFFFFF"))
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.size = Pt(9)
            if r == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    for name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.size = Pt(size)
        style.font.bold = True
    # The built-in Title style can carry a theme-colored paragraph border.
    # Remove it so the manuscript title is plain black, as required for upload.
    title_ppr = styles["Title"]._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Terpene effect claims: an evidence-bounded receptor hypothesis map")
    for value in (
        "Perspective article | evidence-based pharmacological perspective",
        "Daniel C. McShan and Susan Trapp",
        "Daniel C. McShan: Terpedia, LLC | Susan Trapp: [AFFILIATION TO CONFIRM]",
        "ORCID: Daniel C. McShan 0000-0003-3880-1711 | Susan Trapp [TO CONFIRM]",
        "Running title: Terpene receptor hypotheses",
        "Text pages: 11 | Tables: 1 | Figures: 0 | References: 13",
        "Abstract: 153 words | Introduction: 118 words | Discussion: 1,009 words",
        "Abbreviations: CB1, CB2, D2, E0–E4, GABA-A, nAChR, TRPV1",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run(value)
    doc.add_paragraph()

    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("# ") or line.startswith("**Working manuscript") or line.startswith("**Proposed article type"):
            i += 1
            continue
        if line.startswith("**Keywords:**"):
            p = doc.add_paragraph()
            run = p.add_run("Keywords: ")
            run.bold = True
            inline(p, line.split("**Keywords:**", 1)[1].strip())
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(re.sub(r"^## \d+\.\s*", "", line[3:]), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(re.sub(r"^### \d+\.\d+\s*", "", line[4:]), level=2)
            i += 1
            continue
        if line.startswith("| "):
            table_lines = []
            while i < len(lines) and lines[i].startswith("| "):
                if not re.match(r"^\|\s*-", lines[i]):
                    table_lines.append([x.strip() for x in lines[i].strip("|").split("|")])
                i += 1
            add_table(doc, table_lines)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            inline(p, line[2:])
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        add_body_paragraph(doc, line)
        i += 1

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
